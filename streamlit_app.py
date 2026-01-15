import streamlit as st
import json
from docx import Document
from io import BytesIO

# Page config
st.set_page_config(page_title="Lesson Plan Generator", page_icon="📚", layout="centered")

# Custom CSS for clean purple/grey/black/white design
st.markdown("""
<style>
    /* Main container */
    .main {
        background-color: #ffffff;
    }
    
    /* Title styling */
    .title {
        color: #7C3AED;
        font-size: 2.5rem;
        font-weight: 700;
        text-align: center;
        margin-bottom: 1rem;
    }
    
    /* Subtitle */
    .subtitle {
        color: #6B7280;
        text-align: center;
        font-size: 1rem;
        margin-bottom: 2rem;
    }
    
    /* File uploader styling */
    .stFileUploader {
        background-color: #F9FAFB;
        border: 2px dashed #D1D5DB;
        border-radius: 8px;
        padding: 1.5rem;
    }
    
    /* Button styling */
    .stDownloadButton > button {
        background-color: #7C3AED;
        color: white;
        border: none;
        border-radius: 8px;
        padding: 0.75rem 2rem;
        font-weight: 600;
        width: 100%;
        box-shadow: 0 4px 6px rgba(124, 58, 237, 0.2);
        transition: all 0.3s ease;
    }
    
    .stDownloadButton > button:hover {
        background-color: #6D28D9;
        box-shadow: 0 6px 12px rgba(124, 58, 237, 0.3);
        transform: translateY(-2px);
    }
    
    /* Success/Error messages */
    .stSuccess {
        background-color: #F0FDF4;
        border-left: 4px solid #7C3AED;
        color: #374151;
    }
    
    .stError {
        background-color: #FEF2F2;
        border-left: 4px solid #6B7280;
        color: #374151;
    }
    
    .stInfo {
        background-color: #F3F4F6;
        border-left: 4px solid #9CA3AF;
        color: #374151;
    }
    
    /* JSON display */
    .stJson {
        background-color: #1F2937;
        border-radius: 8px;
        padding: 1rem;
    }
    
    /* Expander */
    .streamlit-expanderHeader {
        background-color: #F9FAFB;
        border-radius: 6px;
        color: #374151;
        font-weight: 500;
    }
</style>
""", unsafe_allow_html=True)

# Title
st.markdown('<div class="title">Weekly Lesson Plan Generator</div>', unsafe_allow_html=True)
st.markdown('<div class="subtitle">Paste your JSON data to populate the template</div>', unsafe_allow_html=True)

# JSON text area
st.markdown("### Paste JSON Data")
json_input = st.text_area(
    label="JSON Input",
    height=300,
    placeholder='{\n  "Teacher": "John Doe",\n  "YearClass": "Year 5",\n  "Subject": "English",\n  ...\n}',
    label_visibility="collapsed"
)

# Generate button
if st.button("Generate Lesson Plan", use_container_width=True):
    if json_input:
        try:
            # Parse JSON data
            json_data = json.loads(json_input)
            st.success("✅ JSON data loaded successfully!")
            
            # Show preview
            with st.expander("📄 View Parsed Data"):
                st.json(json_data)
            
            # Load the WLPT.docx template from the templates folder
            try:
                doc = Document("templates/WLPT.docx")
            except FileNotFoundError:
                st.error("❌ WLPT.docx template not found in templates folder.")
                st.stop()
            
            # Function to replace placeholders
            def replace_placeholder(doc, placeholder, value):
                """Replace placeholder text in all tables and paragraphs"""
                # Replace in paragraphs
                for paragraph in doc.paragraphs:
                    if placeholder in paragraph.text:
                        for run in paragraph.runs:
                            if placeholder in run.text:
                                run.text = run.text.replace(placeholder, str(value))
                
                # Replace in tables
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if placeholder in cell.text:
                                for paragraph in cell.paragraphs:
                                    for run in paragraph.runs:
                                        if placeholder in run.text:
                                            run.text = run.text.replace(placeholder, str(value))
            
            # Replace all placeholders with JSON data
            for key, value in json_data.items():
                placeholder = "{{" + key + "}}"
                replace_placeholder(doc, placeholder, value if value else "")
            
            # Save to BytesIO
            docx_buffer = BytesIO()
            doc.save(docx_buffer)
            docx_buffer.seek(0)
            
            st.success("✅ Template populated successfully!")
            
            # Download button
            st.download_button(
                label="📥 Download Lesson Plan",
                data=docx_buffer,
                file_name="Weekly_Lesson_Plan.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )
            
        except json.JSONDecodeError as e:
            st.error(f"❌ Invalid JSON format: {str(e)}")
        except Exception as e:
            st.error(f"❌ Error: {str(e)}")
    else:
        st.info("👆 Please paste your JSON data above")

# Footer with minimal instructions
st.markdown("---")
st.markdown("""
<div style="text-align: center; color: #9CA3AF; font-size: 0.875rem;">
    <p>Use JSON keys like <code>Class1_LearningObjective</code>, <code>Class1_SuccessCriteria</code>, <code>Class1_Vocabulary</code>, etc.</p>
    <p>Include all fields from Class1 through Class5 for each day</p>
</div>
""", unsafe_allow_html=True)
